'''
----------------------------------------------------------------------------------------------------------------------------------------
Author: Sherif Mehralivand
Email: sherif.mehralivand@ukdd.de
Github: https://github.com/smehralivand
Twitter: @smehralivand
Date: 11/04/2024
----------------------------------------------------------------------------------------------------------------------------------------
'''

# Import libraries
from pathlib import Path
from tqdm import tqdm
import win32com.client
from docx import Document

def to_raw(string):
    '''
    Helper function which returns a raw string
    '''
    return r'{}'.format(string)

def docx_to_txt(source_path, target_path):
    docx_path = Path(source_path)
    txt_path = Path(target_path)
    
    # Extract docx files in list
    docx_files = list(docx_path.glob('*.docx'))
    
    num = 0
    for docx_file in tqdm(docx_files):
        document = Document(docx_file)
        content = [p.text for p in document.paragraphs]
        txt_file = docx_file.with_suffix('.txt')
        txt_file =txt_path / txt_file.name
        with open(txt_file, "w", encoding="utf-8") as file:
            file.write("\n".join(content))
        num += 1
    
    return num

def txt_to_docx(source_path, target_path):
    txt_path = Path(source_path)
    docx_path = Path(target_path)
    
    # Extract txt files in list
    txt_files = list(txt_path.glob('*.txt'))

    document = Document()
    num = 0
    for txt_file in tqdm(txt_files):
        with open(txt_file, 'r', encoding='utf-8') as file:
            paragraphs = file.read().split('\n')
        for paragraph in paragraphs:
            document.add_paragraph(paragraph)
        
        docx_file = txt_file.with_suffix('.docx')
        docx_file = docx_path / docx_file.name
        document.save(docx_file)
        num += 1

    return num

def doc_to_docx(source_path, target_path):
    doc_path = Path(source_path)
    docx_path = Path(target_path)
    
    # Define Word Application and make invisible
    word = win32com.client.Dispatch('Word.Application')
    word.visible = False

    # Extract doc files in list
    doc_files = list(doc_path.glob('*.doc'))

    # Iterate over doc files and save as docx format   
    num = 0
    for doc_file in tqdm(doc_files):
        # Convert doc to docx in file name
        docx_file = doc_file.with_suffix('.docx')
        docx_file = docx_path / docx_file.name

        # Convert to absolute paths as raw strings
        doc_file = to_raw(doc_file.absolute())
        docx_file = to_raw(docx_file.absolute())

        # Open file in Word
        wb = word.Documents.Open(doc_file)

        # Save file in different format and close file
        wb.SaveAs2(docx_file, FileFormat=16) # file format for docx
        num += 1
        wb.Close()
    
    # Close Word
    word.Quit()
    
    return num

def docx_to_doc(source_path, target_path):
    docx_path = Path(source_path)
    doc_path = Path(target_path)
    
    # Define Word Application and make invisible
    word = win32com.client.Dispatch('Word.Application')
    word.visible = False

    # Extract docx files in list
    docx_files = list(docx_path.glob('*.docx'))

    # Iterate over docx files and save as doc format   
    num = 0
    for doc_file in tqdm(docx_files):
        # Convert docx to doc in file name
        docx_file = doc_file.with_suffix('.doc')
        docx_file = doc_path / docx_file.name

        # Convert to absolute paths as raw strings
        doc_file = to_raw(doc_file.absolute())
        docx_file = to_raw(docx_file.absolute())

        # Open file in Word
        wb = word.Documents.Open(doc_file)

        # Save file in different format and close file
        wb.SaveAs2(docx_file, FileFormat=0) # file format for doc
        num += 1
        wb.Close()
    
    # Close Word
    word.Quit()
    
    return num
















